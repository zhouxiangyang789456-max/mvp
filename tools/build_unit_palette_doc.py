from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"D:\prounity\mvp\output\docs\单位视觉配色方案.docx")

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(dxa))
    tc_w.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(total))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for edge, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn('w:' + edge))
        if node is None:
            node = OxmlElement('w:' + edge)
            tc_mar.append(node)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')

def font(run, size=10.5, color='172333', bold=False, name='Microsoft YaHei'):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold

def add_palette_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    widths = [1900, 1100, 1900, 4460]
    headers = ['用途 / 部位', '色块', 'HEX', '调整说明']
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, 'E8EEF5')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(text), 10, '1F4D78', True)
    for label, hex_color, note in rows:
        cells = table.add_row().cells
        font(cells[0].paragraphs[0].add_run(label), 10)
        set_cell_shading(cells[1], hex_color)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(cells[2].paragraphs[0].add_run('#' + hex_color), 10, '172333', True, 'Consolas')
        font(cells[3].paragraphs[0].add_run(note), 9.5, '34465A')
    set_table_geometry(table, widths)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Microsoft YaHei'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string('172333')
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2
for style_name, size, color, before, after in [
    ('Heading 1', 16, '2E74B5', 16, 8),
    ('Heading 2', 13, '2E74B5', 12, 6),
]:
    s = styles[style_name]
    s.font.name = 'Microsoft YaHei'
    s._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = True
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(header.add_run('《指挥大师》美术调整参考'), 8.5, '6B7785')
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run('单位视觉配色方案'), 8.5, '7A8794')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)
font(p.add_run('单位视觉配色方案'), 25, '203748', True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
font(p.add_run('士兵与坦克在黄绿色等距地图中的辨识度调整参考'), 12.5, '527083')

callout = doc.add_table(rows=1, cols=1)
callout.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(callout, [9360])
cell = callout.cell(0, 0)
set_cell_shading(cell, 'F1F6FA')
set_cell_margins(cell, 130, 180, 130, 180)
p = cell.paragraphs[0]
font(p.add_run('核心方向：'), 10.5, '1F4D78', True)
font(p.add_run('使用冷蓝灰单位、深蓝黑轮廓与少量阵营色点缀，避免当前棕色单位与森林、沙地及黄绿色平原融为一体。'), 10.5, '34465A')

doc.add_heading('1. 通用轮廓与明暗层级', level=1)
add_palette_table(doc, [
    ('最外轮廓', '172333', '所有外缘使用，替代纯黑；建议保留 2–3 像素。'),
    ('背光轮廓', '26384C', '用于背光侧和轮廓内部的次级分割。'),
    ('最深阴影', '34465A', '加强四肢、炮塔、履带与车身之间的体块区分。'),
    ('受光边缘', 'B8D0DF', '只用于朝光边缘，避免整片高亮。'),
])

doc.add_heading('2. 士兵配色', level=1)
add_palette_table(doc, [
    ('主制服', '3F6077', '覆盖主要布料面积，与黄绿地面形成冷暖对比。'),
    ('制服亮面', '66869A', '用于肩部、胸口和大腿迎光面。'),
    ('制服阴影', '293E52', '用于手臂内侧、胯部和腿部背光面。'),
    ('防弹衣', '59636C', '保持中性灰，与蓝色制服明显分层。'),
    ('头盔', '718696', '比防弹衣稍亮，让头部轮廓更快被识别。'),
    ('枪械', '202D39', '枪管、枪托和瞄具统一使用深色，并加一条亮边。'),
    ('皮肤', 'D1A07D', '面部和手部基础色。'),
    ('皮肤阴影', '916650', '避免皮肤区域与制服亮面混在一起。'),
])

doc.add_page_break()
doc.add_heading('3. 坦克配色', level=1)
add_palette_table(doc, [
    ('主装甲', '48677A', '车身主要面积，保持冷蓝灰基调。'),
    ('装甲亮面', '7896A5', '炮塔顶部、车体上沿和迎光斜面。'),
    ('装甲阴影', '2D4658', '炮塔下缘、侧装甲与车体连接处。'),
    ('履带', '202A32', '必须明显深于车身，形成清楚的底盘轮廓。'),
    ('履带亮边', '53616A', '只画在履带上沿和外缘，表现结构而非整片提亮。'),
    ('炮管', '314757', '比主装甲稍深，轮廓线保持连续。'),
    ('舱盖与接缝', '182A37', '强调炮塔、舱盖、车身面板的结构线。'),
])

doc.add_heading('4. 阵营与编队点缀', level=1)
add_palette_table(doc, [
    ('玩家主色', '2F9FE8', '用于肩甲、头盔条带、坦克侧面识别条。'),
    ('玩家高光', '72D4FF', '只用于玩家识别条的受光边缘。'),
    ('敌人主色', 'D94B45', '用于敌方小面积识别标记。'),
    ('敌人高光', 'FF8A62', '只用于敌方标记的受光边缘。'),
    ('当前激活编队', 'F2C94C', '用于选中环、底光或指挥官标记，不要覆盖单位全身。'),
])

doc.add_heading('5. 血条调整', level=1)
add_palette_table(doc, [
    ('生命值', '45B86B', '替代当前荧光绿，保持可见但不压过单位。'),
    ('血条背景', '172333', '与单位外轮廓统一。'),
    ('血条边框', '0C141D', '提供稳定边界，避免与地形亮部粘连。'),
])

doc.add_heading('6. 执行要点', level=1)
for text in [
    '士兵：头盔、防弹衣、制服和枪械至少保持四个不同的明度层级。',
    '坦克：炮塔比车身略亮，履带明显更暗，炮管保留连续亮边。',
    '阵营色只占单位可见面积的 5%–12%，不要把整台坦克或整套制服染成阵营色。',
    '血条高度建议缩小约 25%，并确保单位轮廓比血条更先被注意到。',
    '最终在实际战斗缩放下检查，而不是只在建模或贴图近景中判断。',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run(text), 10.5)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
